import docx
import sys
import io

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document(r"d:\FPT\Hihi\tailieu\PRO2192_Report_Goc.docx")

for idx, table in enumerate(doc.tables):
    # Check if table has images
    has_img = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "xml" in p._element.xml and "pic:" in p._element.xml:
                    has_img = True
    print(f"Table[{idx}]: {len(table.rows)} rows, {len(table.columns)} columns, has_img: {has_img}")
    if has_img:
        for r_idx, row in enumerate(table.rows):
            print(f"  Row[{r_idx}]: {[cell.text.strip().replace('\n', ' ')[:30] for cell in row.cells]}")
