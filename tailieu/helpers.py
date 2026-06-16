"""
helpers.py — Các hàm tiện ích dùng chung cho toàn bộ tài liệu VisualizationDSA
Áp dụng đúng chuẩn định dạng: Font Arial, kích thước đúng, table header #1F4E79
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Hằng số định dạng ────────────────────────────────────────────────────────
FONT_NAME       = 'Arial'
SIZE_COVER      = 24        # pt — Tiêu đề trang bìa
SIZE_H1         = 18        # pt — Tên Phần (Heading 1)
SIZE_H2         = 14        # pt — Mục con (Heading 2)
SIZE_H3         = 12        # pt — Mục con cấp 3 (in đậm)
SIZE_BODY       = 12        # pt — Văn bản thường
SIZE_CAPTION    = 10        # pt — Chú thích hình/bảng
SIZE_CODE       = 10        # pt — Code block / PlantUML

COLOR_HEADER_BG = '1F4E79'  # Hex — Nền header bảng (xanh đậm FPT)
COLOR_WHITE     = 'FFFFFF'
COLOR_GRAY_ROW  = 'F2F2F2'  # Nền dòng xen kẽ nhẹ
COLOR_BORDER    = 'BFBFBF'  # Viền ô bảng màu xám mờ

# ─── Hàm thiết lập Document ───────────────────────────────────────────────────
def create_document() -> Document:
    """Khởi tạo Document mới với cấu hình trang A4 chuẩn."""
    doc = Document()
    # Thiết lập kích thước trang A4 và lề
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Thiết lập style Normal mặc định
    style = doc.styles['Normal']
    font  = style.font
    font.name = FONT_NAME
    font.size = Pt(SIZE_BODY)

    return doc


def load_document(filepath: str) -> Document:
    """Mở Document đã tồn tại để append thêm nội dung."""
    return Document(filepath)


# ─── Hàm thêm đoạn văn ────────────────────────────────────────────────────────
def add_paragraph(doc: Document, text: str, bold: bool = False,
                  italic: bool = False, size: int = SIZE_BODY,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_after: float = 6.0,
                  line_spacing: float = 1.15) -> None:
    """Thêm đoạn văn bản thường với định dạng chuẩn."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

    run = p.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def add_heading1(doc: Document, text: str) -> None:
    """Thêm tiêu đề Heading 1 (Phần/Chương) — 18pt Bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after  = Pt(10)
    run = p.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(SIZE_H1)
    run.font.bold  = True
    # Đặt màu đen đậm tiêu đề
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_heading2(doc: Document, text: str) -> None:
    """Thêm tiêu đề Heading 2 (Mục con) — 14pt Bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(SIZE_H2)
    run.font.bold  = True


def add_heading3(doc: Document, text: str) -> None:
    """Thêm tiêu đề Heading 3 — 12pt Bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(SIZE_H3)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def add_body(doc: Document, text: str, bold: bool = False,
             italic: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Thêm đoạn văn body căn đều 2 bên — 12pt, spacing 1.15."""
    add_paragraph(doc, text, bold=bold, italic=italic, size=SIZE_BODY,
                  align=align, space_after=6.0, line_spacing=1.15)


def add_caption(doc: Document, text: str) -> None:
    """Thêm chú thích hình hoặc bảng — 10pt Italic, căn giữa."""
    add_paragraph(doc, text, italic=True, size=SIZE_CAPTION,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10.0)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Thêm dòng bullet point (danh sách gạch đầu dòng)."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(20 * (level + 1))
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(SIZE_BODY)


def add_page_break(doc: Document) -> None:
    """Thêm ngắt trang."""
    doc.add_page_break()


def add_image_placeholder(doc: Document, tag: str, caption: str) -> None:
    """
    Thêm placeholder cho hình ảnh (sẽ được thay thế bằng ảnh thực tế sau).
    tag   : Tên placeholder, ví dụ: '[IMAGE_PLACEHOLDER_ERD_OVERVIEW]'
    caption: Chú thích hình ảnh.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(2)
    run = p.add_run(tag)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)  # Đỏ để dễ nhận ra
    add_caption(doc, caption)


def add_image(doc: Document, image_path: str, caption: str, width_cm: float = 15.0) -> None:
    """
    Chèn ảnh thực tế vào tài liệu từ đường dẫn file.
    Nếu file ảnh không tồn tại, tự động chuyển sang chèn placeholder để tránh crash code.
    """
    if not os.path.exists(image_path):
        # Fallback sang placeholder nếu chưa có file ảnh thực tế
        tag_name = f'[IMAGE_PLACEHOLDER: {os.path.basename(image_path)}]'
        add_image_placeholder(doc, tag_name, caption)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(4)
    
    # Chèn ảnh thực tế
    try:
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
        add_caption(doc, caption)
    except Exception as e:
        # Nếu có lỗi chèn (ví dụ định dạng file ảnh hỏng)
        tag_name = f'[ERROR_LOADING_IMAGE: {os.path.basename(image_path)} - {str(e)}]'
        add_image_placeholder(doc, tag_name, caption)



def add_code_block(doc: Document, code: str, lang: str = 'PlantUML') -> None:
    """
    Thêm khối code (PlantUML / Mermaid / SQL) — font Courier New, nền xám nhạt.
    """
    # Nếu là code mô tả diagram (UML, Mermaid, Sitemap, Wireframe) thì không đưa vào Word
    # mà ghi ra file hinhanh_uml_source.txt để lưu trữ.
    uml_langs = {'PlantUML', 'Mermaid', 'Mermaid ERD', 'Sitemap Tree', 'Wireframe ASCII', 'Text Architecture Diagram'}
    lang_lower = lang.lower()
    if lang in uml_langs or 'uml' in lang_lower or 'mermaid' in lang_lower or 'sitemap' in lang_lower or 'wireframe' in lang_lower:
        txt_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'hinhanh_uml_source.txt')
        try:
            with open(txt_path, 'a', encoding='utf-8') as f:
                f.write(f'{"="*80}\n')
                f.write(f'DIAGRAM SOURCE: {lang}\n')
                f.write(f'{"="*80}\n')
                f.write(code)
                f.write('\n\n')
        except Exception as e:
            print(f'Lỗi khi ghi diagram source: {e}')
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    pf.left_indent  = Pt(12)

    # Thêm label ngôn ngữ
    lbl = doc.add_paragraph()
    lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_lbl = lbl.add_run(f'[{lang}]')
    run_lbl.font.name  = FONT_NAME
    run_lbl.font.size  = Pt(9)
    run_lbl.font.bold  = True
    run_lbl.font.italic = True
    run_lbl.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Thêm code text
    run_code = p.add_run(code)
    run_code.font.name = 'Courier New'
    run_code.font.size = Pt(SIZE_CODE)

    # Đặt nền xám nhạt cho đoạn code
    _set_paragraph_shading(p, 'F5F5F5')


def _set_paragraph_shading(paragraph, color_hex: str) -> None:
    """Đặt màu nền cho paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


# ─── Hàm tạo và định dạng Bảng ───────────────────────────────────────────────
def set_cell_bg(cell, color_hex: str) -> None:
    """Đặt màu nền cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)


def set_cell_border(cell, color_hex: str = COLOR_BORDER) -> None:
    """Đặt viền cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def format_header_cell(cell, text: str, font_size: int = 11) -> None:
    """
    Định dạng ô header bảng:
    - Nền #1F4E79 (xanh đậm)
    - Chữ trắng, Bold, căn giữa
    """
    set_cell_bg(cell, COLOR_HEADER_BG)
    set_cell_border(cell, COLOR_HEADER_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Xóa nội dung cũ và thêm mới
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(font_size)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def format_data_cell(cell, text: str, center: bool = False,
                      bold: bool = False, font_size: int = 11,
                      alt_row: bool = False) -> None:
    """
    Định dạng ô dữ liệu bảng:
    - Căn trái (hoặc căn giữa nếu center=True)
    - Có viền xám mờ
    - Nền xen kẽ nhẹ nếu alt_row=True
    """
    set_cell_border(cell, COLOR_BORDER)
    if alt_row:
        set_cell_bg(cell, COLOR_GRAY_ROW)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)

    for run in para.runs:
        run.text = ''

    run = para.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = Pt(font_size)
    run.font.bold  = bold


def create_table(doc: Document, headers: list,
                 col_widths: list = None) -> object:
    """
    Tạo bảng mới với hàng header được định dạng chuẩn.
    headers   : Danh sách tên cột header.
    col_widths: Danh sách độ rộng cột (Inches). Mặc định tự động.
    Trả về đối tượng table để tiếp tục thêm dòng dữ liệu.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Đặt độ rộng cột
    if col_widths:
        for i, (cell, w) in enumerate(zip(table.columns, col_widths)):
            for c in cell.cells:
                c.width = Inches(w)

    # Định dạng header row
    hdr_row = table.rows[0]
    for i, (cell, hdr_text) in enumerate(zip(hdr_row.cells, headers)):
        format_header_cell(cell, hdr_text)

    return table


def add_table_row(table, values: list, centers: list = None,
                  bolds: list = None, alt_row: bool = False,
                  font_size: int = 11) -> None:
    """
    Thêm một dòng dữ liệu vào bảng.
    values  : Danh sách giá trị text cho từng ô.
    centers : Danh sách bool xác định ô nào căn giữa (mặc định False).
    bolds   : Danh sách bool xác định ô nào in đậm.
    """
    row = table.add_row()
    if centers is None:
        centers = [False] * len(values)
    if bolds is None:
        bolds = [False] * len(values)
    for cell, val, ctr, bld in zip(row.cells, values, centers, bolds):
        format_data_cell(cell, str(val), center=ctr, bold=bld,
                         font_size=font_size, alt_row=alt_row)


# ─── Hàm tiện ích khác ───────────────────────────────────────────────────────
def add_section_divider(doc: Document) -> None:
    """Thêm đường kẻ ngang phân cách."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)


def get_output_path() -> str:
    """Trả về đường dẫn file output cố định."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, 'PRO2192_VisualizationDSA_Report.docx')


def save_document(doc: Document, filepath: str) -> None:
    """Lưu document và thông báo."""
    doc.save(filepath)
    print(f'[OK] Đã lưu: {filepath}')
