import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import sys
import io

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Đường dẫn
DOC_GOC = r"d:\FPT\Hihi\tailieu\PRO2192_Report_Goc.docx"
DOC_OUT = r"d:\FPT\Hihi\document\PRO2192_Report.docx"
IMG_DIR = r"d:\FPT\Hihi\tailieu\images"
UML_OUT = r"d:\FPT\Hihi\tailieu\hinhanh_uml_source.txt"

# Đảm bảo thư mục output tồn tại
os.makedirs(os.path.dirname(DOC_OUT), exist_ok=True)

# Khởi tạo hoặc dọn dẹp file UML source
with open(UML_OUT, "w", encoding="utf-8") as f:
    f.write("========================================================\n")
    f.write("DANH SÁCH MÃ NGUỒN BIỂU ĐỒ UML / MERMAID - VISUALIZATIONDSA\n")
    f.write("========================================================\n\n")

# Map thông tin cũ -> mới
old_to_new = {
    # Tên thành viên
    "Nguyễn Đình Thiên Long": "Mai Tiểu Bảo",
    "Nguyễn Duy Phương": "Thái Quang Sơn",
    "Võ Hoàng Tú": "Huỳnh Lê Minh Thư",
    "Trần Minh Hiếu": "Trần Viết Tâm Phúc",
    "Phạm Minh Hậu": "",
    "Trần Vũ Quang Huy": "",
    # Mã sinh viên
    "PS12345": "TD01287",
    "PS12346": "TD01282",
    "PS12347": "TD01131",
    "PS12348": "TD01261",
    "PS12349": "",
    "PS12350": "",
    # Viết tắt tên trong bảng Sprint
    "N.Đ.T Long": "M.T Bảo",
    "N.D Phương": "T.Q Sơn",
    "V.H Tú": "H.L.M Thư",
    "T.M Hiếu": "T.V.T Phúc",
    "P.M Hậu": "H.L.M Thư",
    "T.V.Q Huy": "T.V.T Phúc",
}

def replace_text_safely(p, mapping):
    text = p.text
    has_change = False
    for old_val, new_val in mapping.items():
        if old_val in text:
            text = text.replace(old_val, new_val)
            has_change = True
    if has_change:
        # Giữ style và font của run đầu tiên
        font_name = "Times New Roman"
        font_size = Pt(12)
        bold = False
        italic = False
        color_rgb = None
        
        if p.runs:
            first_run = p.runs[0]
            if first_run.font.name: font_name = first_run.font.name
            if first_run.font.size: font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic
            color_rgb = first_run.font.color.rgb
            
        p.text = ""
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = font_size
        if bold is not None: r.font.bold = bold
        if italic is not None: r.font.italic = italic
        if color_rgb: r.font.color.rgb = color_rgb

def clean_html_tags(p):
    """Lọc bỏ các thẻ HTML như <p>, <i> khỏi chú thích ảnh và định dạng italic."""
    text = p.text.strip()
    if text.startswith("<p><i>") and text.endswith("</i></p>"):
        cleaned_text = text.replace("<p><i>", "").replace("</i></p>", "")
        p.text = ""
        r = p.add_run(cleaned_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

print("📖 Đang đọc file mẫu gốc...")
doc = docx.Document(DOC_GOC)

# 1. Thay thế thông tin cơ bản trong toàn bộ paragraph (ngoại trừ các placeholder ảnh/UML)
print("✍️ Đang cập nhật thông tin thành viên trong các đoạn văn...")
for p in doc.paragraphs:
    txt = p.text.strip()
    # Bỏ qua các placeholder hình ảnh/UML để xử lý riêng sau
    if any(k in txt for k in ["[KHUNG CHÈN ẢNH", "[SƠ ĐỒ MERMAID", "[IMAGE_PLACEHOLDER"]):
        continue
    replace_text_safely(p, old_to_new)

# 2. Xử lý các bảng biểu
print("📊 Đang cập nhật các bảng biểu...")
# Bảng 0: Danh sách thành viên rút gọn ở trang phụ bìa
table_0 = doc.tables[0]
# Cập nhật thông tin nhóm mới
# Hàng 2: Bảo
table_0.cell(2, 1).text = "Mai Tiểu Bảo"
table_0.cell(2, 2).text = "Nhóm trưởng (Backend & Core Engines)"
# Hàng 3: Sơn
table_0.cell(3, 1).text = "Thái Quang Sơn"
table_0.cell(3, 2).text = "Thành viên nòng cốt (Frontend & Visualizations)"
# Hàng 4: Thư
table_0.cell(4, 1).text = "Huỳnh Lê Minh Thư"
table_0.cell(4, 2).text = "Thành viên (Database & API Services)"
# Hàng 5: Phúc
table_0.cell(5, 1).text = "Trần Viết Tâm Phúc"
table_0.cell(5, 2).text = "Thành viên (DevOps & Gamification)"
# Xóa hàng 6 và hàng 7 thừa
table_0._tbl.remove(table_0.rows[7]._tr)
table_0._tbl.remove(table_0.rows[6]._tr)

# Định dạng lại font cho Table 0
for row in table_0.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

# Bảng 1: Bảng phân công nhiệm vụ chi tiết
table_1 = doc.tables[1]
# Cập nhật thông tin chi tiết
members_tasks = [
    ("Mai Tiểu Bảo", "TD01287", "Nhóm trưởng", "Thiết kế kiến trúc hệ thống ASP.NET Core Clean Architecture, Database Schema PostgreSQL, Authentication JWT, Security Middleware và điều phối dự án."),
    ("Thái Quang Sơn", "TD01282", "Thành viên", "Lập trình lõi render đồ họa Canvas 2D (60 FPS), nội suy Lerp hoạt ảnh, thiết kế VCR Playback Timeline, Monaco IDE Integration và SOLID/OOP Sandboxes."),
    ("Huỳnh Lê Minh Thư", "TD01131", "Thành viên", "Viết tài liệu đặc tả yêu cầu hệ thống (SRS), lập trình API Endpoints, thuật toán Strategy (Sort/Graph/Tree), và viết bộ kịch bản kiểm thử tích hợp (Vitest/xUnit)."),
    ("Trần Viết Tâm Phúc", "TD01261", "Thành viên", "Thiết lập môi trường Docker container, cấu hình triển khai CI/CD VPS, lập trình hệ thống Quiz trắc nghiệm, Leaderboard và viết tài liệu HDSD.")
]
for idx, (name, mssv, role, task) in enumerate(members_tasks, start=1):
    table_1.cell(idx, 0).text = name
    table_1.cell(idx, 1).text = mssv
    table_1.cell(idx, 2).text = role
    table_1.cell(idx, 3).text = task

# Xóa hàng 5 và hàng 6 thừa
table_1._tbl.remove(table_1.rows[6]._tr)
table_1._tbl.remove(table_1.rows[5]._tr)

# Định dạng Table 1
for r_idx, row in enumerate(table_1.rows):
    for cell in row.cells:
        for p in cell.paragraphs:
            # Canh lề trái cho cột nhiệm vụ, canh giữa cho cột khác
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if p.alignment is None else p.alignment
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                if r_idx == 0:
                    r.font.bold = True

# Bảng 2: Kế hoạch Sprint
table_2 = doc.tables[2]
for row in table_2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            replace_text_safely(p, old_to_new)

# Các bảng DB & API
for t_idx in range(3, len(doc.tables)):
    for row in doc.tables[t_idx].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)

# 3. Chèn ảnh chân dung Ban dự án dạng bảng Borderless
print("👥 Đang tạo bảng Ban dự án không viền...")

# Tìm vị trí bắt đầu chèn ảnh Ban dự án
# P[61] là '1.2 Ban dự án'. Ta sẽ xóa P[64] -> P[71] và chèn bảng Ban dự án vào vị trí này.
para_indices_to_delete = []
target_insert_p_idx = -1

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if i >= 64 and i <= 71:
        para_indices_to_delete.append(i)
        if i == 64:
            target_insert_p_idx = i

# Lấy paragraph để chèn bảng trước nó
insert_before_p = doc.paragraphs[target_insert_p_idx]

# Tạo bảng Ban dự án mới: 4 hàng, 2 cột
new_table = doc.add_table(rows=4, cols=2)
# Di chuyển bảng lên vị trí mong muốn
insert_before_p._element.addprevious(new_table._element)

# Thiết lập bảng borderless
tblPr = new_table._tbl.tblPr
borders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
tblPr.append(borders)

# Điền thông tin 4 thành viên mới
new_members_info = [
    ("MAI TIỂU BẢO", "avatar_bao.png", [
        "MSSV: TD01287 — Nhóm trưởng (Backend & Core)",
        "Thiết kế và triển khai kiến trúc Clean Architecture cho backend C# .NET",
        "Lập trình Web API, cấu hình Dependency Injection và Middleware",
        "Thiết kế sơ đồ cơ sở dữ liệu PostgreSQL và phân quyền Auth (JWT)",
        "Quản trị dự án, điều phối và phân chia công việc trong 12 Sprints"
    ]),
    ("THÁI QUANG SƠN", "avatar_son.png", [
        "MSSV: TD01282 — Thành viên nòng cốt (Frontend & Vis)",
        "Lập trình lõi render đồ họa Canvas 2D (60 FPS) và nội suy Lerp hoạt ảnh",
        "Lập trình các cấu trúc dữ liệu Sorting, Graph, Tree và VCR Playback controls",
        "Xây dựng sân chơi vẽ đồ thị tự do (Force-directed layout)",
        "Mô phỏng các khái niệm OOP VTable, SOLID Sandboxes và DI Container"
    ]),
    ("HUỲNH LÊ MINH THƯ", "avatar_thu.png", [
        "MSSV: TD01131 — Thành viên (Backend & QA)",
        "Thiết kế và biên soạn tài liệu đặc tả yêu cầu hệ thống (SRS)",
        "Lập trình các thuật toán Strategy C# (Bubble/Quick/Dijkstra/DFS)",
        "Viết các kịch bản kiểm thử tích hợp (xUnit/Vitest) đảm bảo độ phủ code",
        "Phát triển giao diện Dashboard học viên và phân hệ slide bài giảng"
    ]),
    ("TRẦN VIẾT TÂM PHÚC", "avatar_phuc.png", [
        "MSSV: TD01261 — Thành viên (DevOps & Testing)",
        "Đóng gói container Docker, cấu hình build và Deploy lên VPS Cloud",
        "Lập trình hệ thống câu hỏi trắc nghiệm Quiz Card, checkpoint bài học",
        "Xây dựng tính năng Gamification tích lũy XP, thăng hạng Leaderboard",
        "Viết tài liệu hướng dẫn sử dụng và báo cáo kiểm thử chất lượng (QA)"
    ])
]

for idx, (name, img_name, tasks) in enumerate(new_members_info):
    row = new_table.rows[idx]
    # Cột 0: Chèn hình ảnh
    cell_img = row.cells[0]
    cell_img.width = Cm(3.5)
    p_img = cell_img.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img = p_img.add_run()
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        r_img.add_picture(img_path, width=Cm(3.0))
    
    # Cột 1: Chèn văn bản
    cell_txt = row.cells[1]
    cell_txt.width = Cm(12.5)
    
    # Tên thành viên
    p_name = cell_txt.paragraphs[0]
    p_name.paragraph_format.space_before = Pt(4)
    p_name.paragraph_format.space_after = Pt(4)
    r_name = p_name.add_run(name)
    r_name.font.name = "Times New Roman"
    r_name.font.size = Pt(13)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(180, 0, 0) # Màu đỏ nâu học thuật
    
    # Danh sách nhiệm vụ
    for task_str in tasks:
        p_task = cell_txt.add_paragraph()
        p_task.paragraph_format.space_after = Pt(2)
        p_task.paragraph_format.line_spacing = 1.15
        
        # Format bullet point
        if task_str == tasks[0]:
            p_task.paragraph_format.left_indent = Cm(0.2)
            r_task = p_task.add_run(task_str)
            r_task.font.bold = True
        else:
            p_task.paragraph_format.left_indent = Cm(0.5)
            r_task = p_task.add_run("• " + task_str)
            
        r_task.font.name = "Times New Roman"
        r_task.font.size = Pt(11)

# Xóa các paragraph placeholder cũ
# Lưu ý xóa ngược từ dưới lên để tránh lệch chỉ số
for i in sorted(para_indices_to_delete, reverse=True):
    p_to_del = doc.paragraphs[i]
    p_to_del._element.getparent().remove(p_to_del._element)

# 4. Quét và chèn hình ảnh giao diện & sơ đồ thực tế
print("🖼️ Đang quét và chèn các hình ảnh thực tế...")

# Map placeholder ảnh -> ảnh thực tế
image_mappings = {
    "deployment_diagram.png": "client_server.png",
    "sitemap_diagram.png": "sitemap.png",
    "layout_mockup.png": "wireframe_main.png",
    "ui_auth.png": "screen_landing.png",
    "ui_dashboard.png": "screen_dashboard.png",
    "ui_visualizer.png": "screen_bubble_sort.png",
    "ui_comparison.png": "screen_dijkstra.png",
    "ui_graph_playground.png": "screen_dijkstra.png",
    "ui_monaco_ide.png": "screen_bubble_sort.png",
    "ui_slide_quiz.png": "screen_quiz.png",
    "ui_concurrency.png": "screen_oop.png",
    "ui_solid_sandbox.png": "screen_solid.png",
    "ui_embed_widget.png": "screen_teacher.png",
    "ui_leaderboard.png": "screen_dashboard.png",
    # Các diagram UML/Mermaid
    "class_clean_arch": "class_clean_arch.png",
    "seq_algorithm": "seq_algorithm.png",
    "activity_quiz": "activity_quiz.png",
    "class_strategy": "class_strategy.png"
}

# Quét tất cả paragraph để xử lý chèn ảnh và dọn dẹp UML code
p_idx = 0
while p_idx < len(doc.paragraphs):
    p = doc.paragraphs[p_idx]
    txt = p.text.strip()
    
    # A. Xử lý placeholder ảnh bình thường
    if "[KHUNG CHÈN ẢNH THỰC TẾ:" in txt:
        # Tìm file ảnh thực tế
        img_file_found = None
        for key, val in image_mappings.items():
            if key in txt:
                img_file_found = val
                break
        
        if img_file_found:
            img_path = os.path.join(IMG_DIR, img_file_found)
            p.text = "" # Xóa text placeholder
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_img = p.add_run()
            if os.path.exists(img_path):
                # Set width cho ảnh to (14.5 cm là tỷ lệ tối ưu trên trang A4)
                r_img.add_picture(img_path, width=Cm(14.5))
                print(f"  -> Đã chèn ảnh: {img_file_found}")
            else:
                p.text = f"[HÌNH ẢNH THIẾU: {img_file_found}]"
        else:
            p.text = "[Khung ảnh trống]"
            
    # B. Xử lý placeholder UML / Mermaid code
    elif "[SƠ ĐỒ MERMAID" in txt or "classDiagram" in txt or "sequenceDiagram" in txt:
        # Ghi mã nguồn UML vào file hinhanh_uml_source.txt
        # Ghi nguồn UML
        uml_title = "UML DIAGRAM"
        img_file = "client_server.png" # Mặc định
        
        if "classDiagram" in txt:
            if "Application" in txt:
                uml_title = "Class Diagram: Cấu trúc ứng dụng Client-side"
                img_file = "class_clean_arch.png"
            else:
                uml_title = "Class Diagram: Core Engine"
                img_file = "class_strategy.png"
        elif "sequenceDiagram" in txt:
            uml_title = "Sequence Diagram: Biên dịch & Chạy mã nguồn giải thuật"
            img_file = "seq_algorithm.png"
        elif "stateDiagram" in txt:
            uml_title = "State Diagram: Vòng đời hoạt cảnh VCR"
            img_file = "activity_quiz.png"
        elif "Student((Học Viên))" in txt or "Student" in txt:
            uml_title = "Use Case Diagram: Học viên"
            img_file = "uc_student.png"
        elif "Admin((Quản Trị Viên))" in txt or "Admin" in txt:
            uml_title = "Use Case Diagram: Quản trị viên"
            img_file = "uc_admin.png"
        elif "graph TD" in txt:
            uml_title = "Use Case Diagram"
            img_file = "uc_student.png"
        
        # Ghi file
        with open(UML_OUT, "a", encoding="utf-8") as f:
            f.write(f"### {uml_title} ###\n")
            f.write(txt + "\n\n")
            f.write("-" * 50 + "\n\n")
            
        # Thay thế paragraph UML bằng hình ảnh thực tế tương ứng
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join(IMG_DIR, img_file)
        r_img = p.add_run()
        if os.path.exists(img_path):
            r_img.add_picture(img_path, width=Cm(14.5))
            print(f"  -> Đã chèn sơ đồ UML: {img_file}")
        else:
            p.text = f"[UML DIAGRAM: {uml_title}]"
            
    # C. Dọn dẹp thẻ HTML chú thích ảnh
    elif txt.startswith("<p><i>") and txt.endswith("</i></p>"):
        clean_html_tags(p)
        
    p_idx += 1

# Dọn dẹp thêm các paragraph chú thích ảnh sau khi chèn ảnh
for p in doc.paragraphs:
    if p.text.strip().startswith("<p><i>"):
        clean_html_tags(p)

print("💾 Đang lưu file báo cáo hoàn thiện...")
doc.save(DOC_OUT)
print("✅ Hoàn thành cập nhật báo cáo tốt nghiệp thành công tại: document/PRO2192_Report.docx")
