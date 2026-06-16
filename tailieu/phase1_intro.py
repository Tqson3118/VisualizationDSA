"""
phase1_intro.py
Phần 1: Giới thiệu đề tài — Tổng quan dự án + Ban dự án (4 thành viên)
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers import *
from docx.shared import Pt, Inches


def build_phan1(doc: Document) -> None:

    add_heading1(doc, 'PHẦN 1: GIỚI THIỆU ĐỀ TÀI')
    add_section_divider(doc)

    # ── 1.1 Giới thiệu dự án ──────────────────────────────────────────────────
    add_heading2(doc, '1.1. Giới thiệu dự án')

    add_image(doc, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'images', 'logo_project.png'),
              'Hình 1.1: Logo dự án VisualizationDSA', width_cm=8.0)

    add_body(doc,
        'VISUALIZATION DSA là một nền tảng giáo dục kỹ thuật số (EdTech Platform) '
        'thế hệ mới, được thiết kế chuyên biệt để trực quan hóa toàn bộ hệ sinh thái '
        'Khoa học Máy tính — từ những thuật toán sắp xếp căn bản nhất như Bubble Sort, '
        'Quick Sort cho đến các cấu trúc đồ thị phức tạp như Dijkstra, BFS/DFS, cây '
        'BST, và thậm chí cả các khái niệm phần mềm trừu tượng bậc cao như OOP, '
        'nguyên tắc SOLID, Design Patterns và DI Container.')

    add_body(doc,
        'Điểm khác biệt cốt lõi của hệ thống so với các công cụ tương tự trên thị '
        'trường (VisuAlgo, CS Visualizations...) là kiến trúc Data-Driven Animation: '
        'Backend C# đóng vai trò "não bộ" — chạy thuật toán trong bộ nhớ, bắt giữ '
        'từng trạng thái thay đổi (State Frames) và sinh ra mảng JSON tối ưu hóa. '
        'Frontend Vue 3 đóng vai trò "màn hình chiếu" — chỉ đơn thuần đọc dữ liệu '
        'JSON và phát lại (Playback) thành hoạt hình Canvas 60 FPS mượt mà, không '
        'bao giờ bị giật lag hay sai lệch logic thuật toán.')

    add_heading3(doc, '1.1.1. Tính năng nổi bật')
    features = [
        ('DSA Visualization Engine',
         'Trực quan hóa 15+ thuật toán: Bubble Sort, Quick Sort, Merge Sort, '
         'Heap Sort, Binary Search, BFS, DFS, Dijkstra, BST, Stack, Queue...'),
        ('OOP & SOLID Sandbox',
         'Mô phỏng trực quan các tính chất OOP (Kế thừa, Đa hình, Đóng gói) '
         'và 5 nguyên tắc SOLID bằng các mô hình hạt khói và kết nối Bezier.'),
        ('Design Patterns Visualizer',
         'Minh họa động các mẫu thiết kế phổ biến: Observer, Factory, Strategy, '
         'Decorator thông qua hoạt ảnh Canvas tương tác.'),
        ('DI Container Simulator',
         'Trực quan hóa quá trình Dependency Injection — xem cách các Services '
         'được tiêm (Inject) vào nhau ở Runtime.'),
        ('Interactive Quiz System',
         'Hệ thống trắc nghiệm tương tác 30+ câu hỏi liên kết với từng module '
         'bài học, tích hợp chấm điểm tự động và trao thưởng XP.'),
        ('Gamification Engine',
         'Hệ thống tích lũy điểm kinh nghiệm XP, thăng cấp bậc (8 Level), '
         'mở khóa huy hiệu thành tựu và bảng xếp hạng (Leaderboard).'),
        ('E-Lecture Mode',
         'Bài giảng kịch bản dẫn dắt từng bước — người học không bị "ném" '
         'vào màn hình code trống mà được hướng dẫn có cấu trúc.'),
        ('Algorithm Comparison',
         'So sánh hiệu năng thực tế (thời gian chạy, số bước) giữa tối đa '
         '4 thuật toán cùng một lúc trên cùng một bộ dữ liệu đầu vào.'),
        ('Interactive Playground',
         'Sân chơi tự do — người dùng tự vẽ Node, Edge, xây dựng đồ thị '
         'của riêng mình và để thuật toán chạy trực tiếp trên đó.'),
        ('Premium & Payment',
         'Hệ thống thanh toán nâng cấp Premium với tích hợp cổng thanh toán, '
         'mở khóa toàn bộ nội dung độc quyền cao cấp.'),
    ]

    table = create_table(doc,
                         ['STT', 'Tính năng', 'Mô tả chi tiết'],
                         col_widths=[0.5, 1.6, 4.8])
    for i, (feat, desc) in enumerate(features):
        add_table_row(table, [str(i + 1), feat, desc],
                      centers=[True, False, False],
                      bolds=[False, True, False],
                      alt_row=(i % 2 == 1))

    doc.add_paragraph()

    add_heading3(doc, '1.1.2. Mục tiêu dự án')
    goals = [
        'Xây dựng nền tảng trực quan hóa đa chiều, bao phủ từ DSA căn bản đến '
        'kiến trúc phần mềm nâng cao (OOP, SOLID, Design Patterns, DI).',
        'Thiết kế hệ thống animation engine đạt 60 FPS ổn định, '
        'sử dụng kỹ thuật Vector Lerp và Double Buffering Canvas.',
        'Xây dựng Backend RESTful API theo chuẩn Clean Architecture với '
        'Strategy Pattern, đảm bảo khả năng mở rộng (Extensibility) tối đa.',
        'Tích hợp hệ thống Gamification và Quiz để tăng động lực học tập, '
        'đo lường tiến trình và khen thưởng người học.',
        'Triển khai hệ thống phân quyền đa vai trò (Student, Teacher, Admin) '
        'với bảo mật JWT và khả năng quản lý người dùng toàn diện.',
    ]
    for goal in goals:
        add_bullet(doc, goal)

    doc.add_paragraph()

    add_heading3(doc, '1.1.3. Phạm vi và giới hạn')
    add_body(doc,
        'Dự án tập trung vào đối tượng chính là sinh viên ngành Công nghệ Thông '
        'tin và những người tự học lập trình. Hệ thống hỗ trợ đầy đủ trên trình '
        'duyệt web hiện đại (Chrome, Firefox, Edge) trên thiết bị máy tính. '
        'Phiên bản hiện tại chưa tối ưu hoàn toàn cho thiết bị di động màn hình '
        'nhỏ và chưa tích hợp chế độ ngoại tuyến (Offline mode).')

    # ── 1.2 Ban dự án ─────────────────────────────────────────────────────────
    add_heading2(doc, '1.2. Ban dự án')

    add_body(doc,
        'Dự án được thực hiện bởi nhóm 4 sinh viên chuyên ngành Lập trình Ứng dụng '
        'phần mềm, Trường Cao đẳng FPT Polytechnic. Nhóm vận hành theo mô hình '
        'Agile Scrum với các Sprint 2 tuần, phân chia công việc theo năng lực '
        'và chuyên môn của từng thành viên.')

    # Bảng thành viên
    table_members = create_table(doc,
        ['STT', 'Họ và tên', 'MSSV', 'Vai trò trong nhóm', 'Chuyên trách chính'],
        col_widths=[0.4, 1.5, 0.9, 1.4, 2.7])

    members_data = [
        ('1', 'Huỳnh Lê Minh Thư',  'TD01131', 'Thành viên',
         'Khảo sát yêu cầu, xây dựng Vue.js Components giao diện (Landing, Dashboard, Profile), '
         'viết tài liệu SRS và kiểm thử giao diện người dùng (UI Testing)'),
        ('2', 'Trần Viết Tâm Phúc', 'TD01261', 'Thành viên',
         'Khảo sát thị trường và phân tích đối thủ cạnh tranh, xây dựng Vue.js Components '
         'phụ trợ (Quiz UI, Leaderboard, Payment UI), viết tài liệu kiểm thử và kịch bản test'),
        ('3', 'Mai Tiểu Bảo',        'TD01287', 'Nhóm trưởng',
         'Thiết kế kiến trúc Clean Architecture, xây dựng Domain Layer & EF Core '
         'Database Schema, triển khai Algorithm Strategy Classes (BFS, DFS, Dijkstra, BST), '
         'thiết kế API RESTful, xây dựng Admin Panel và hệ thống phân quyền'),
        ('4', 'Thái Quang Sơn',      'TD01282', 'Thành viên nòng cốt',
         'Xây dựng Animation Engine (60 FPS Canvas), toàn bộ Frontend DSA Modules '
         '(Sorting, Graph, OOP, SOLID, DI Visualization), thiết kế hệ thống Gamification '
         'và Quiz, tích hợp JWT Auth, StatelessAuth Strategy và hệ thống Payment'),
    ]

    for i, row in enumerate(members_data):
        add_table_row(table_members, list(row),
                      centers=[True, False, True, False, False],
                      bolds=[False, True, False, True, False],
                      alt_row=(i % 2 == 1))

    doc.add_paragraph()
    add_caption(doc, 'Bảng 1.1: Danh sách thành viên nhóm và vai trò')
    
    add_image(doc, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'images', 'team_members.png'),
              'Hình 1.2: Ảnh đại diện nhóm ban dự án', width_cm=14.0)

    # Chi tiết từng thành viên
    add_heading3(doc, 'Chi tiết đóng góp theo thành viên')

    detail_members = [
        {
            'name': '👤 Mai Tiểu Bảo (TD01287) — Nhóm trưởng',
            'items': [
                'Thiết kế và triển khai kiến trúc Clean Architecture 4 tầng cho backend C# .NET 10',
                'Xây dựng toàn bộ Domain Entities: User, Quiz, QuizAttempt, Badge, Order, LearningProgress',
                'Triển khai EF Core DbContext, Migrations và Database Schema PostgreSQL',
                'Phát triển 10+ Algorithm Strategy: BFSStrategy, DFSStrategy, DijkstraStrategy, BSTStrategy, BinarySearchStrategy, QueueStrategy, StackStrategy, LinearSearchStrategy, MergeSortStrategy, HeapSortStrategy...',
                'Thiết kế Admin Panel (AdminController, AdminPanelView.vue) với đầy đủ CRUD quản lý',
                'Xây dựng hệ thống phân quyền RBAC (Student, Teacher, Admin) với JWT Bearer',
                'Điều phối Sprint Planning, Review và phân công công việc hàng tuần',
            ]
        },
        {
            'name': '👤 Thái Quang Sơn (TD01282) — Kỹ sư nòng cốt',
            'items': [
                'Xây dựng toàn bộ Canvas Animation Engine (60 FPS, Vector Lerp, Double Buffering)',
                'Phát triển 15+ DSA Module Visualizers: BubbleSortRenderer, QuickSortRenderer, BarChartRenderer...',
                'Triển khai OOPConceptsStrategy, SOLIDPrinciplesStrategy, DIContainerStrategy, DesignPatternsStrategy',
                'Xây dựng hệ thống Gamification hoàn chỉnh: XP, Level, Streak, Badge (GamificationStrategy)',
                'Phát triển StatelessAuthStrategy với in-memory fallback và JWT Chained Token',
                'Tích hợp StatelessPaymentStrategy và cổng thanh toán Premium',
                'Xây dựng hệ thống Quiz hoàn chỉnh: StatelessQuizController + QuizBankStrategy (30+ câu hỏi)',
                'Xây dựng Algorithm Comparison feature (CompareController, CompareView)',
            ]
        },
        {
            'name': '👤 Huỳnh Lê Minh Thư (TD01131) — Frontend & Documentation',
            'items': [
                'Thiết kế và xây dựng Landing Page, Dashboard View, Profile View',
                'Phát triển các Vue.js Components phụ trợ: Header, Sidebar, Navigation',
                'Viết tài liệu yêu cầu hệ thống (SRS), Use Case, User Story',
                'Thực hiện khảo sát người dùng và phân tích yêu cầu từ góc độ UX',
                'Kiểm thử giao diện người dùng (UI/UX Testing) các màn hình chính',
                'Biên soạn hướng dẫn sử dụng hệ thống',
            ]
        },
        {
            'name': '👤 Trần Viết Tâm Phúc (TD01261) — Testing & Survey',
            'items': [
                'Khảo sát và phân tích các hệ thống tương tự: VisuAlgo, CS Visualizations, AlgoExpert',
                'Xây dựng Quiz UI Components và Leaderboard View',
                'Phát triển Payment UI (PremiumCheckoutView.vue)',
                'Viết toàn bộ kịch bản kiểm thử (Test Cases) cho các module cốt lõi',
                'Thực hiện kiểm thử tích hợp API và kiểm thử end-to-end',
                'Tổng hợp và soạn thảo báo cáo tốt nghiệp',
            ]
        },
    ]

    for member in detail_members:
        add_heading3(doc, member['name'])
        for item in member['items']:
            add_bullet(doc, item)
        doc.add_paragraph()

    add_page_break(doc)


def run() -> None:
    filepath = get_output_path()
    doc = load_document(filepath)
    print('[Phase 1] Đang thêm: Phần 1 – Giới thiệu đề tài...')
    build_phan1(doc)
    save_document(doc, filepath)
    print('[Phase 1] HOÀN THÀNH.\n')


if __name__ == '__main__':
    run()
