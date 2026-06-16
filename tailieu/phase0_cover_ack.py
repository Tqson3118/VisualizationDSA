"""
phase0_cover_ack.py
Trang bìa + Phụ bìa + Lời cảm ơn + Mục lục + Danh mục Hình + Danh mục Bảng + Lời mở đầu
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers import *
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_cover(doc: Document) -> None:
    """Trang bìa chính — theo chuẩn FPT Polytechnic."""
    # --- Tiêu đề trường ---
    p_school = doc.add_paragraph()
    p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_school.paragraph_format.space_before = Pt(0)
    p_school.paragraph_format.space_after  = Pt(2)
    r = p_school.add_run('TRƯỜNG CAO ĐẲNG FPT POLYTECHNIC')
    r.font.name = FONT_NAME; r.font.size = Pt(13); r.font.bold = True

    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.space_after = Pt(0)
    r2 = p_dept.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
    r2.font.name = FONT_NAME; r2.font.size = Pt(13); r2.font.bold = True

    add_section_divider(doc)

    # --- Logo FPT Polytechnic ---
    add_image(doc, os.path.join(os.path.dirname(os.path.abspath(__file__)), 'images', 'logo_fpt.png'),
              'Logo FPT Polytechnic', width_cm=8.0)

    # --- Nhãn loại tài liệu ---
    for _ in range(2):
        doc.add_paragraph()

    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = p_label.add_run('BÁO CÁO DỰ ÁN TỐT NGHIỆP')
    r_lbl.font.name = FONT_NAME; r_lbl.font.size = Pt(16); r_lbl.font.bold = True

    p_code = doc.add_paragraph()
    p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_code = p_code.add_run('Mã học phần: PRO2192')
    r_code.font.name = FONT_NAME; r_code.font.size = Pt(13)
    p_code.paragraph_format.space_after = Pt(20)

    # --- Tên đề tài ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(4)
    r_t = p_title.add_run('ĐỀ TÀI:')
    r_t.font.name = FONT_NAME; r_t.font.size = Pt(14); r_t.font.bold = True

    p_project = doc.add_paragraph()
    p_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_project.paragraph_format.space_after = Pt(4)
    r_proj = p_project.add_run('VISUALIZATION DSA')
    r_proj.font.name = FONT_NAME; r_proj.font.size = Pt(SIZE_COVER); r_proj.font.bold = True
    r_proj.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    r_sub = p_sub.add_run('Nền tảng trực quan hóa Cấu trúc dữ liệu và Giải thuật')
    r_sub.font.name = FONT_NAME; r_sub.font.size = Pt(14); r_sub.font.italic = True

    add_section_divider(doc)

    # --- Thông tin nhóm ---
    for _ in range(1):
        doc.add_paragraph()

    members = [
        ('GIẢNG VIÊN HƯỚNG DẪN', '...............................................'),
        ('NHÓM SINH VIÊN THỰC HIỆN', ''),
        ('1. Huỳnh Lê Minh Thư',   'MSSV: TD01131'),
        ('2. Trần Viết Tâm Phúc',  'MSSV: TD01261'),
        ('3. Mai Tiểu Bảo',         'MSSV: TD01287  (Nhóm trưởng)'),
        ('4. Thái Quang Sơn',       'MSSV: TD01282'),
        ('CHUYÊN NGÀNH',            'Lập trình Ứng dụng phần mềm'),
        ('NIÊN KHÓA',               '2023 – 2026'),
    ]
    for label, value in members:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_m.paragraph_format.left_indent  = Cm(3)
        p_m.paragraph_format.space_after  = Pt(4)
        r_l = p_m.add_run(f'{label}: ')
        r_l.font.name = FONT_NAME; r_l.font.size = Pt(12); r_l.font.bold = True
        r_v = p_m.add_run(value)
        r_v.font.name = FONT_NAME; r_v.font.size = Pt(12)

    # --- Thời gian & địa điểm ---
    for _ in range(2):
        doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run('Hà Nội / Hồ Chí Minh / Đà Nẵng, năm 2026')
    r_date.font.name = FONT_NAME; r_date.font.size = Pt(12); r_date.font.italic = True

    add_page_break(doc)


def build_phụ_bia(doc: Document) -> None:
    """Phụ bìa nội dung (trang thứ hai sau trang bìa)."""
    add_heading1(doc, 'XÁC NHẬN CỦA GIẢNG VIÊN HƯỚNG DẪN')

    add_body(doc,
        'Tôi xác nhận rằng nhóm sinh viên thực hiện đề tài "Visualization DSA — Nền tảng '
        'trực quan hóa Cấu trúc dữ liệu và Giải thuật" đã hoàn thành đầy đủ các yêu cầu '
        'kỹ thuật, tài liệu mô tả và báo cáo theo đúng quy chuẩn của môn học PRO2192.')

    add_body(doc, 'Hà Nội, ngày ........ tháng ........ năm 2026')

    for _ in range(3):
        doc.add_paragraph()

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sig = p_sig.add_run('Giảng viên hướng dẫn\n(Ký và ghi rõ họ tên)')
    r_sig.font.name = FONT_NAME; r_sig.font.size = Pt(12); r_sig.font.bold = True

    add_page_break(doc)


def build_loi_cam_on(doc: Document) -> None:
    """Lời cảm ơn."""
    add_heading1(doc, 'LỜI CẢM ƠN')
    add_section_divider(doc)

    add_body(doc,
        'Trước tiên, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc nhất '
        'đến quý thầy cô giáo Khoa Công nghệ Thông tin, Trường Cao đẳng FPT Polytechnic '
        'đã tận tâm hướng dẫn, truyền đạt kiến thức chuyên ngành và kinh nghiệm thực tiễn '
        'trong suốt quá trình học tập tại trường.')

    add_body(doc,
        'Đặc biệt, chúng em xin gửi lời cảm ơn sâu sắc đến Giảng viên hướng dẫn đã '
        'dành nhiều thời gian quý báu để chỉ dẫn, góp ý và định hướng cho nhóm trong '
        'quá trình nghiên cứu, phân tích và phát triển dự án tốt nghiệp. Những nhận xét '
        'và phản hồi của thầy/cô là nguồn động lực lớn nhất để nhóm hoàn thiện sản phẩm.')

    add_body(doc,
        'Dự án "Visualization DSA" được xây dựng với mục tiêu biến những khái niệm '
        'trừu tượng của Cấu trúc dữ liệu, Giải thuật và Nguyên tắc thiết kế phần mềm '
        '(OOP, SOLID, Design Patterns) thành những hoạt hình tương tác trực quan, giúp '
        'người học dễ dàng tiếp cận và nắm bắt bản chất vấn đề. Hành trình phát triển '
        'hệ thống này đã giúp nhóm tích lũy được nhiều kinh nghiệm thực tế quý giá về '
        'kiến trúc Clean Architecture (.NET 10), lập trình Vue 3 Composition API, '
        'thiết kế cơ sở dữ liệu PostgreSQL và quy trình phát triển phần mềm Agile Scrum.')

    add_body(doc,
        'Chúng em cũng xin gửi lời cảm ơn đến gia đình, bạn bè và các anh chị đi '
        'trước đã chia sẻ tài liệu, hỗ trợ tinh thần và tạo điều kiện thuận lợi nhất '
        'để nhóm hoàn thành dự án tốt nghiệp này.')

    add_body(doc,
        'Mặc dù nhóm đã cố gắng hết sức, song do thời gian và kinh nghiệm còn hạn chế, '
        'báo cáo chắc chắn không tránh khỏi những thiếu sót. Nhóm rất mong nhận được '
        'sự góp ý chân thành từ quý thầy cô và hội đồng để tiếp tục hoàn thiện.')

    for _ in range(2):
        doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run(
        'Hà Nội, tháng 06 năm 2026\n'
        'Nhóm tác giả\n\n'
        'Huỳnh Lê Minh Thư  –  Trần Viết Tâm Phúc\n'
        'Mai Tiểu Bảo  –  Thái Quang Sơn'
    )
    r_s.font.name = FONT_NAME; r_s.font.size = Pt(12); r_s.font.italic = True

    add_page_break(doc)


def build_muc_luc(doc: Document) -> None:
    """Mục lục (tĩnh — người dùng cập nhật trường TOC trong Word)."""
    add_heading1(doc, 'MỤC LỤC')
    add_section_divider(doc)

    toc_entries = [
        ('LỜI CẢM ƠN', ''),
        ('MỤC LỤC', ''),
        ('DANH MỤC HÌNH ẢNH', ''),
        ('DANH MỤC BẢNG BIỂU', ''),
        ('LỜI MỞ ĐẦU', ''),
        ('PHẦN 1: GIỚI THIỆU ĐỀ TÀI', ''),
        ('    1.1. Giới thiệu dự án', ''),
        ('    1.2. Ban dự án', ''),
        ('PHẦN 2: KHẢO SÁT – SURVEY', ''),
        ('    2.1. Yêu cầu của khách hàng', ''),
        ('    2.2. Kế hoạch dự án', ''),
        ('PHẦN 3: PHÂN TÍCH – ANALYSIS', ''),
        ('    3.1. Mô hình triển khai hệ thống', ''),
        ('    3.2. Sơ đồ Use Cases', ''),
        ('    3.3. Đặc tả yêu cầu hệ thống (SRS)', ''),
        ('PHẦN 4: THIẾT KẾ – DESIGN', ''),
        ('    4.1. Mô hình công nghệ', ''),
        ('    4.2. Thiết kế giao diện', ''),
        ('    4.3. Thiết kế dữ liệu (ERD)', ''),
        ('    4.4. Sơ đồ lớp (Class Diagram)', ''),
        ('PHẦN 5: THỰC HIỆN – IMPLEMENT', ''),
        ('    5.1. Database – Cơ sở dữ liệu', ''),
        ('    5.2. Layout thực tế', ''),
        ('    5.3. Sơ đồ kiến trúc công nghệ thực tế', ''),
        ('    5.4. Các loại sơ đồ (Sequence, Activity, Class)', ''),
        ('    5.5. API – Danh sách Controllers và Services', ''),
        ('PHẦN 6: KIỂM THỬ – TESTING', ''),
        ('PHẦN 7: ĐÓNG GÓI & TRIỂN KHAI – DEPLOYMENT', ''),
        ('KẾT LUẬN & HƯỚNG PHÁT TRIỂN', ''),
        ('TÀI LIỆU THAM KHẢO', ''),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        is_main = not entry.startswith('    ')
        run = p.add_run(entry)
        run.font.name = FONT_NAME
        run.font.size = Pt(12 if is_main else 11)
        run.font.bold = is_main
        if is_main:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_page_break(doc)


def build_danh_muc_hinh(doc: Document) -> None:
    """Danh mục hình ảnh."""
    add_heading1(doc, 'DANH MỤC HÌNH ẢNH')
    add_section_divider(doc)

    figures = [
        ('Hình 1.1',  'Logo dự án VisualizationDSA'),
        ('Hình 1.2',  'Giao diện trang chủ (Landing Page)'),
        ('Hình 1.3',  'Màn hình Dashboard học viên'),
        ('Hình 3.1',  'Mô hình Client-Server tổng thể'),
        ('Hình 3.2',  'Sơ đồ Use Case – Guest (Khách vãng lai)'),
        ('Hình 3.3',  'Sơ đồ Use Case – Student (Học viên)'),
        ('Hình 3.4',  'Sơ đồ Use Case – Teacher (Giảng viên)'),
        ('Hình 3.5',  'Sơ đồ Use Case – Admin (Quản trị viên)'),
        ('Hình 4.1',  'Mô hình công nghệ tổng thể (Tech Stack)'),
        ('Hình 4.2',  'Sitemap toàn bộ ứng dụng'),
        ('Hình 4.3',  'Wireframe trang Visualization chính'),
        ('Hình 4.4',  'Sơ đồ ERD tổng quan cơ sở dữ liệu'),
        ('Hình 4.5',  'Sơ đồ ERD chi tiết bảng users và user_progress'),
        ('Hình 4.6',  'Sơ đồ lớp – Kiến trúc Clean Architecture'),
        ('Hình 4.7',  'Sơ đồ lớp – Strategy Pattern cho Algorithm'),
        ('Hình 5.1',  'Kết quả migration PostgreSQL thực tế'),
        ('Hình 5.2',  'Giao diện thực tế – Bubble Sort Visualization'),
        ('Hình 5.3',  'Giao diện thực tế – Graph Dijkstra Visualization'),
        ('Hình 5.4',  'Giao diện thực tế – OOP Sandbox'),
        ('Hình 5.5',  'Giao diện thực tế – SOLID Principles Viz'),
        ('Hình 5.6',  'Giao diện thực tế – Admin Panel'),
        ('Hình 5.7',  'Giao diện thực tế – Teacher Panel'),
        ('Hình 5.8',  'Sơ đồ kiến trúc Clean Architecture 4 tầng'),
        ('Hình 5.9',  'Sequence Diagram – Luồng thực thi thuật toán'),
        ('Hình 5.10', 'Sequence Diagram – Luồng xác thực JWT'),
        ('Hình 5.11', 'Activity Diagram – Luồng làm Quiz và nhận XP'),
        ('Hình 5.12', 'Class Diagram – Strategy Pattern backend C#'),
        ('Hình 6.1',  'Kết quả kiểm thử API với Swagger UI'),
        ('Hình 7.1',  'Quy trình build và deploy ứng dụng'),
    ]

    table = create_table(doc, ['STT', 'Tên hình', 'Trang'],
                         col_widths=[0.6, 5.5, 0.8])
    for i, (fig_id, fig_name) in enumerate(figures):
        add_table_row(table,
                      [str(i + 1), f'{fig_id}: {fig_name}', '...'],
                      centers=[True, False, True],
                      alt_row=(i % 2 == 1))

    doc.add_paragraph()
    add_page_break(doc)


def build_danh_muc_bang(doc: Document) -> None:
    """Danh mục bảng biểu."""
    add_heading1(doc, 'DANH MỤC BẢNG BIỂU')
    add_section_divider(doc)

    tables = [
        ('Bảng 1.1',  'Danh sách thành viên nhóm và vai trò'),
        ('Bảng 2.1',  'Danh sách yêu cầu chức năng (Functional Requirements)'),
        ('Bảng 2.2',  'Danh sách yêu cầu phi chức năng (Non-Functional Requirements)'),
        ('Bảng 2.3',  'Kế hoạch dự án – Phân chia Sprint và công việc'),
        ('Bảng 2.4',  'Phân công chi tiết công việc theo từng thành viên'),
        ('Bảng 3.1',  'Đặc tả Use Case UC-01: Xem trực quan hóa thuật toán'),
        ('Bảng 3.2',  'Đặc tả Use Case UC-02: Đăng ký tài khoản'),
        ('Bảng 3.3',  'Đặc tả Use Case UC-03: Làm Quiz và nhận XP'),
        ('Bảng 3.4',  'Đặc tả Use Case UC-04: Quản lý người dùng (Admin)'),
        ('Bảng 4.1',  'Danh sách công nghệ sử dụng trong dự án'),
        ('Bảng 4.2',  'Đặc tả bảng users (Thông tin tài khoản người dùng)'),
        ('Bảng 4.3',  'Đặc tả bảng algorithms (Thư viện thuật toán)'),
        ('Bảng 4.4',  'Đặc tả bảng quizzes (Bài kiểm tra trắc nghiệm)'),
        ('Bảng 4.5',  'Đặc tả bảng quiz_questions (Câu hỏi chi tiết)'),
        ('Bảng 4.6',  'Đặc tả bảng user_submissions (Lịch sử nộp bài)'),
        ('Bảng 4.7',  'Đặc tả bảng achievements (Thành tựu và huy hiệu)'),
        ('Bảng 5.1',  'Danh sách 22 API Controllers và chức năng'),
        ('Bảng 5.2',  'Danh sách 29 Algorithm Strategy Classes'),
        ('Bảng 5.3',  'Bảng tổng hợp API Endpoints chính'),
        ('Bảng 6.1',  'Kịch bản kiểm thử – Module Authentication'),
        ('Bảng 6.2',  'Kịch bản kiểm thử – Module Algorithm Visualization'),
        ('Bảng 6.3',  'Kịch bản kiểm thử – Module Quiz System'),
        ('Bảng 6.4',  'Kịch bản kiểm thử – Module Admin Panel'),
        ('Bảng 7.1',  'Hướng dẫn cấu hình môi trường triển khai'),
    ]

    table = create_table(doc, ['STT', 'Tên bảng', 'Trang'],
                         col_widths=[0.6, 5.5, 0.8])
    for i, (tbl_id, tbl_name) in enumerate(tables):
        add_table_row(table,
                      [str(i + 1), f'{tbl_id}: {tbl_name}', '...'],
                      centers=[True, False, True],
                      alt_row=(i % 2 == 1))

    doc.add_paragraph()
    add_page_break(doc)


def build_loi_mo_dau(doc: Document) -> None:
    """Lời mở đầu."""
    add_heading1(doc, 'LỜI MỞ ĐẦU')
    add_section_divider(doc)

    add_body(doc,
        'Trong thời đại Cách mạng Công nghiệp 4.0, lĩnh vực Công nghệ Thông tin đang '
        'chứng kiến tốc độ phát triển chưa từng có. Cùng với sự bùng nổ đó, nhu cầu '
        'đào tạo lập trình viên có tư duy thuật toán vững chắc và khả năng thiết kế '
        'hệ thống phần mềm chất lượng cao ngày càng trở nên cấp thiết hơn bao giờ hết.')

    add_body(doc,
        'Thực tế giảng dạy và học tập cho thấy, phần lớn sinh viên ngành Công nghệ '
        'Thông tin gặp khó khăn đáng kể khi tiếp cận các khái niệm trừu tượng như '
        'Cấu trúc Dữ liệu (Data Structures), Giải thuật (Algorithms), hay các Nguyên '
        'tắc thiết kế phần mềm SOLID, mô hình Design Patterns. Nguyên nhân chính là '
        'các tài nguyên học tập hiện tại đa phần chỉ trình bày dưới dạng văn bản và '
        'pseudocode tĩnh, thiếu tính tương tác và khả năng trực quan hóa động.')

    add_body(doc,
        'Xuất phát từ bài toán thực tế đó, nhóm nghiên cứu và phát triển dự án '
        '"VISUALIZATION DSA" — một nền tảng E-Learning tương tác thế hệ mới được '
        'xây dựng nhằm biến các khái niệm phức tạp thành những hoạt hình Canvas 60 FPS '
        'sinh động, trực quan và dễ hiểu. Hệ thống không chỉ đơn thuần là công cụ '
        'minh họa thuật toán, mà còn tích hợp đầy đủ hệ sinh thái học tập hoàn chỉnh '
        'bao gồm: Hệ thống Trắc nghiệm tương tác (Quiz System), Gamification với tích '
        'lũy điểm kinh nghiệm XP và thăng cấp bậc, Sân chơi tự do (Interactive '
        'Playground), Chế độ bài giảng kịch bản (E-Lecture Mode), và Công cụ so sánh '
        'hiệu năng thuật toán (Algorithm Comparison).')

    add_body(doc,
        'Về mặt kỹ thuật, hệ thống được xây dựng theo kiến trúc Client-Server hiện đại '
        'với Backend sử dụng .NET 10 / C# theo chuẩn Clean Architecture và Frontend '
        'sử dụng Vue 3 Composition API + TypeScript + Vite. Đây là công nghệ được '
        'sử dụng rộng rãi trong môi trường doanh nghiệp thực tế, giúp nhóm tích lũy '
        'được nhiều kinh nghiệm có giá trị cho sự nghiệp lập trình sau này.')

    add_body(doc,
        'Báo cáo này trình bày toàn bộ quá trình phân tích, thiết kế và triển khai '
        'hệ thống Visualization DSA theo quy trình phát triển phần mềm chuyên nghiệp, '
        'từ giai đoạn Khảo sát yêu cầu (Survey), Phân tích (Analysis), Thiết kế '
        '(Design), Lập trình (Implementation), Kiểm thử (Testing) cho đến Đóng gói '
        'và Triển khai (Deployment).')

    add_page_break(doc)


def run() -> None:
    """Tạo file document mới và build toàn bộ Phase 0."""
    filepath = get_output_path()
    doc = create_document()
    print('[Phase 0] Đang tạo: Trang bìa, Lời cảm ơn, Mục lục, Danh mục...')
    build_cover(doc)
    build_phụ_bia(doc)
    build_loi_cam_on(doc)
    build_muc_luc(doc)
    build_danh_muc_hinh(doc)
    build_danh_muc_bang(doc)
    build_loi_mo_dau(doc)
    save_document(doc, filepath)
    print('[Phase 0] HOÀN THÀNH.\n')


if __name__ == '__main__':
    run()
