"""
phase2_survey.py
Phần 2: Khảo sát – Yêu cầu khách hàng + Kế hoạch dự án
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers import *
from docx.shared import Pt


def build_phan2(doc: Document) -> None:

    add_heading1(doc, 'PHẦN 2: KHẢO SÁT – SURVEY')
    add_section_divider(doc)

    # ── 2.1 Yêu cầu của khách hàng ────────────────────────────────────────────
    add_heading2(doc, '2.1. Yêu cầu của khách hàng')

    add_body(doc,
        'Thông qua quá trình khảo sát, phỏng vấn giảng viên và sinh viên ngành '
        'Công nghệ Thông tin tại các cơ sở đào tạo, nhóm đã tổng hợp được bộ yêu cầu '
        'đầy đủ và chi tiết cho hệ thống VisualizationDSA.')

    add_heading3(doc, '2.1.1. Yêu cầu chức năng (Functional Requirements)')

    func_reqs = [
        ('FR-01', 'Trực quan hóa thuật toán',
         'Hệ thống phải hiển thị được hoạt hình trực quan cho 15+ thuật toán '
         'sắp xếp, tìm kiếm, đồ thị và cây nhị phân với tốc độ 60 FPS.',
         'Bắt buộc', 'Cao'),
        ('FR-02', 'Điều khiển VCR Playback',
         'Người dùng có thể Play, Pause, Step Forward/Backward, điều chỉnh tốc '
         'độ (0.5x, 1x, 2x) và kéo thanh tiến trình Seekbar.', 'Bắt buộc', 'Cao'),
        ('FR-03', 'Nhập dữ liệu tùy chỉnh',
         'Người dùng có thể tự nhập mảng dữ liệu đầu vào tùy ý, hệ thống tự '
         'động parse, validate và chạy thuật toán trên dữ liệu đó.', 'Bắt buộc', 'Cao'),
        ('FR-04', 'Đăng ký / Đăng nhập',
         'Hỗ trợ đăng ký tài khoản bằng email, đăng nhập bảo mật JWT, '
         'và quản lý phiên đăng nhập. Admin có thể khóa/mở tài khoản.', 'Bắt buộc', 'Cao'),
        ('FR-05', 'Hệ thống Quiz tương tác',
         '30+ câu hỏi trắc nghiệm liên kết với từng module bài học, '
         'chấm điểm tự động, hiển thị đáp án giải thích sau khi nộp bài.', 'Bắt buộc', 'Cao'),
        ('FR-06', 'Gamification (XP & Level)',
         'Tích lũy điểm kinh nghiệm XP khi học, làm quiz, đăng nhập liên tiếp. '
         'Thăng cấp tự động qua 8 Level với huy hiệu thành tựu Neon.', 'Bắt buộc', 'Cao'),
        ('FR-07', 'OOP & SOLID Visualization',
         'Trực quan hóa 3 tính chất OOP và 5 nguyên tắc SOLID bằng '
         'hoạt hình hạt khói, Bezier connector và glass morphism UI.', 'Bắt buộc', 'Cao'),
        ('FR-08', 'Design Patterns Visualizer',
         'Minh họa động Observer, Factory, Strategy, Decorator bằng '
         'Canvas tương tác với chú thích học thuật rõ ràng.', 'Bắt buộc', 'Trung bình'),
        ('FR-09', 'DI Container Simulation',
         'Mô phỏng quá trình Dependency Injection — trực quan hóa '
         'luồng tiêm phụ thuộc vào các Service trong IOC Container.', 'Bắt buộc', 'Trung bình'),
        ('FR-10', 'Leaderboard & Bảng xếp hạng',
         'Hiển thị top 10 học viên có điểm XP cao nhất, '
         'cập nhật thời gian thực khi có người thăng cấp.', 'Mong muốn', 'Trung bình'),
        ('FR-11', 'Algorithm Comparison',
         'So sánh đồng thời tối đa 4 thuật toán trên cùng input, '
         'hiển thị số bước, thời gian thực thi server-side.', 'Mong muốn', 'Trung bình'),
        ('FR-12', 'Interactive Playground',
         'Người dùng tự vẽ Node, Edge bằng chuột (drag & drop), '
         'thuật toán đồ thị chạy trực tiếp trên graph tự vẽ.', 'Mong muốn', 'Cao'),
        ('FR-13', 'E-Lecture Mode',
         'Chế độ bài giảng dẫn dắt từng bước với narration text, '
         'highlight dòng pseudocode tương ứng với hoạt hình.', 'Mong muốn', 'Trung bình'),
        ('FR-14', 'Premium Subscription',
         'Hệ thống thanh toán nâng cấp tài khoản Premium, mở khóa '
         'toàn bộ nội dung độc quyền và tắt quảng cáo.', 'Bắt buộc', 'Cao'),
        ('FR-15', 'Teacher Panel',
         'Giảng viên quản lý Quiz, xem thống kê học tập của học viên, '
         'tạo/sửa/xóa câu hỏi trắc nghiệm và bài giảng.', 'Bắt buộc', 'Cao'),
        ('FR-16', 'Admin Panel',
         'Admin quản lý toàn bộ người dùng: phân quyền, khóa/mở tài khoản, '
         'nâng cấp Premium, xem Dashboard thống kê hệ thống.', 'Bắt buộc', 'Cao'),
    ]

    table_fr = create_table(doc,
        ['Mã YC', 'Tên yêu cầu', 'Mô tả', 'Loại', 'Độ ưu tiên'],
        col_widths=[0.7, 1.4, 3.5, 0.9, 0.9])
    for i, row in enumerate(func_reqs):
        add_table_row(table_fr, list(row),
                      centers=[True, False, False, True, True],
                      bolds=[True, True, False, False, False],
                      alt_row=(i % 2 == 1))
    doc.add_paragraph()
    add_caption(doc, 'Bảng 2.1: Danh sách yêu cầu chức năng (Functional Requirements)')

    add_heading3(doc, '2.1.2. Yêu cầu phi chức năng (Non-Functional Requirements)')

    nonfunc_reqs = [
        ('NFR-01', 'Hiệu năng',
         'Animation Canvas đạt tối thiểu 55 FPS. API phản hồi < 15ms cho '
         'dữ liệu tĩnh, < 50ms cho ghi dữ liệu.'),
        ('NFR-02', 'Bảo mật',
         'Sử dụng JWT Bearer Token, mã hóa mật khẩu BCrypt, '
         'bảo vệ tất cả endpoints yêu cầu xác thực bằng middleware.'),
        ('NFR-03', 'Độ ổn định',
         'Hệ thống hoạt động ổn định 24/7. Có cơ chế fallback in-memory '
         'khi database PostgreSQL tạm thời ngắt kết nối.'),
        ('NFR-04', 'Khả năng mở rộng',
         'Strategy Pattern cho phép thêm thuật toán mới không cần sửa '
         'code hiện tại (tuân thủ Open/Closed Principle).'),
        ('NFR-05', 'Khả năng sử dụng',
         'Người dùng mới có thể bắt đầu xem visualization trong < 2 phút '
         'mà không cần đọc tài liệu hướng dẫn.'),
        ('NFR-06', 'Tương thích',
         'Hỗ trợ đầy đủ Chrome 100+, Firefox 100+, Edge 100+. '
         'Màn hình tối thiểu 1280px để trải nghiệm đầy đủ.'),
        ('NFR-07', 'Bộ nhớ RAM',
         'Mức tiêu thụ RAM máy khách ổn định 15–22MB sau 5 phút sử dụng '
         '(không rò rỉ bộ nhớ từ Canvas particles).'),
        ('NFR-08', 'Thẩm mỹ',
         'Giao diện Glassmorphism + Neon Dark Theme nhất quán. '
         'Bảng màu HSL chuẩn hóa: Cyan (chạy), Amber (cảnh báo), Emerald (đúng).'),
    ]

    table_nfr = create_table(doc,
        ['Mã YC', 'Danh mục', 'Đặc tả yêu cầu phi chức năng'],
        col_widths=[0.8, 1.2, 4.9])
    for i, row in enumerate(nonfunc_reqs):
        add_table_row(table_nfr, list(row),
                      centers=[True, False, False],
                      bolds=[True, True, False],
                      alt_row=(i % 2 == 1))
    doc.add_paragraph()
    add_caption(doc, 'Bảng 2.2: Danh sách yêu cầu phi chức năng (Non-Functional Requirements)')

    # ── 2.2 Kế hoạch dự án ────────────────────────────────────────────────────
    add_heading2(doc, '2.2. Kế hoạch dự án')

    add_body(doc,
        'Dự án VisualizationDSA được phát triển theo phương pháp Agile Scrum '
        'với các Sprint 2 tuần. Toàn bộ lộ trình được chia thành 3 giai đoạn '
        'lớn tương ứng với 7 Sprint chính:')

    sprints = [
        ('Sprint 1', 'Tuần 1–2',
         'Khảo sát, phân tích yêu cầu, thiết kế kiến trúc hệ thống',
         'Thư, Phúc (60%), Bảo, Sơn (40%)'),
        ('Sprint 2', 'Tuần 3–4',
         'Thiết kế ERD, API Spec, cài đặt môi trường, khởi tạo dự án',
         'Bảo, Sơn (70%), Thư, Phúc (30%)'),
        ('Sprint 3', 'Tuần 5–6',
         'Backend: Domain Layer, EF Core, Auth JWT; Frontend: Landing, Dashboard',
         'Bảo, Sơn (75%), Thư, Phúc (25%)'),
        ('Sprint 4', 'Tuần 7–8',
         'DSA Visualization: Sorting, Graph; Animation Engine 60 FPS',
         'Bảo, Sơn (80%), Thư, Phúc (20%)'),
        ('Sprint 5', 'Tuần 9–10',
         'OOP/SOLID/DI Visualization; Quiz System; Gamification Engine',
         'Bảo, Sơn (75%), Thư, Phúc (25%)'),
        ('Sprint 6', 'Tuần 11–12',
         'Admin Panel, Teacher Panel; Payment; Leaderboard; Playground',
         'Bảo, Sơn (70%), Thư, Phúc (30%)'),
        ('Sprint 7', 'Tuần 13–14',
         'Kiểm thử toàn diện, fix bugs, hoàn thiện tài liệu, deploy',
         'Thư, Phúc (60%), Bảo, Sơn (40%)'),
    ]

    table_sprint = create_table(doc,
        ['Sprint', 'Thời gian', 'Nội dung công việc', 'Phân công'],
        col_widths=[0.7, 0.9, 3.8, 1.8])
    for i, row in enumerate(sprints):
        add_table_row(table_sprint, list(row),
                      centers=[True, True, False, False],
                      bolds=[True, False, False, False],
                      alt_row=(i % 2 == 1))
    doc.add_paragraph()
    add_caption(doc, 'Bảng 2.3: Kế hoạch dự án – Phân chia Sprint và công việc')

    add_heading3(doc, '2.2.1. Phân công chi tiết theo công việc')
    add_body(doc,
        'Dựa trên năng lực và thế mạnh của từng thành viên, công việc được phân '
        'chia như sau — đảm bảo các task cốt lõi về code, kiến trúc và phân tích '
        'hệ thống chiếm 65–70% tổng khối lượng thuộc về Sơn và Bảo:')

    tasks = [
        ('T-01', 'Khảo sát yêu cầu & phỏng vấn người dùng',
         'Thư, Phúc', 'Thư, Phúc', '3 ngày', '✅'),
        ('T-02', 'Phân tích đối thủ cạnh tranh (VisuAlgo, AlgoExpert)',
         'Phúc', 'Phúc', '2 ngày', '✅'),
        ('T-03', 'Viết tài liệu SRS (Software Requirements Specification)',
         'Thư', 'Thư', '3 ngày', '✅'),
        ('T-04', 'Thiết kế kiến trúc Clean Architecture backend',
         'Bảo', 'Bảo, Sơn', '2 ngày', '✅'),
        ('T-05', 'Thiết kế ERD & Database Schema PostgreSQL',
         'Bảo', 'Bảo', '2 ngày', '✅'),
        ('T-06', 'Thiết kế RESTful API Specification',
         'Bảo, Sơn', 'Bảo, Sơn', '2 ngày', '✅'),
        ('T-07', 'Triển khai Domain Entities (User, Quiz, Badge, Order...)',
         'Bảo', 'Bảo', '3 ngày', '✅'),
        ('T-08', 'Triển khai EF Core, Migrations, Database Seed',
         'Bảo', 'Bảo', '2 ngày', '✅'),
        ('T-09', 'Xây dựng JWT Auth System & Phân quyền RBAC',
         'Sơn', 'Bảo, Sơn', '3 ngày', '✅'),
        ('T-10', 'Xây dựng Canvas Animation Engine (60 FPS, Lerp)',
         'Sơn', 'Sơn', '5 ngày', '✅'),
        ('T-11', 'Triển khai Sorting Algorithms (Bubble, Quick, Merge, Heap, Radix)',
         'Bảo, Sơn', 'Bảo, Sơn', '4 ngày', '✅'),
        ('T-12', 'Triển khai Graph Algorithms (BFS, DFS, Dijkstra, BST)',
         'Bảo', 'Bảo, Sơn', '4 ngày', '✅'),
        ('T-13', 'Triển khai OOP, SOLID, Design Patterns, DI Visualization',
         'Sơn', 'Sơn', '6 ngày', '✅'),
        ('T-14', 'Xây dựng Quiz System & QuizBankStrategy (30+ câu hỏi)',
         'Sơn', 'Sơn', '4 ngày', '✅'),
        ('T-15', 'Xây dựng Gamification Engine (XP, Level, Badge, Streak)',
         'Sơn', 'Bảo, Sơn', '3 ngày', '✅'),
        ('T-16', 'Xây dựng Admin Panel & Teacher Panel',
         'Bảo, Sơn', 'Bảo, Sơn', '4 ngày', '✅'),
        ('T-17', 'Xây dựng Landing Page, Dashboard, Profile (Vue.js)',
         'Thư', 'Thư, Sơn', '4 ngày', '✅'),
        ('T-18', 'Xây dựng Quiz UI, Leaderboard, Payment UI (Vue.js)',
         'Phúc', 'Phúc, Sơn', '3 ngày', '✅'),
        ('T-19', 'Tích hợp Payment & Premium Subscription',
         'Sơn', 'Bảo, Sơn', '3 ngày', '✅'),
        ('T-20', 'Viết Test Cases & thực hiện kiểm thử',
         'Phúc, Thư', 'Phúc, Thư', '5 ngày', '✅'),
        ('T-21', 'Deploy & cấu hình môi trường production',
         'Bảo, Sơn', 'Bảo, Sơn', '2 ngày', '✅'),
        ('T-22', 'Hoàn thiện tài liệu báo cáo tốt nghiệp',
         'Tất cả', 'Thư, Phúc', '4 ngày', '✅'),
    ]

    table_tasks = create_table(doc,
        ['Mã', 'Tên công việc', 'Phụ trách chính', 'Hỗ trợ', 'Thời gian', 'TT'],
        col_widths=[0.5, 3.0, 1.1, 1.1, 0.8, 0.4])
    for i, row in enumerate(tasks):
        add_table_row(table_tasks, list(row),
                      centers=[True, False, True, True, True, True],
                      bolds=[True, False, False, False, False, False],
                      alt_row=(i % 2 == 1))
    doc.add_paragraph()
    add_caption(doc, 'Bảng 2.4: Phân công chi tiết công việc theo từng thành viên')

    add_body(doc,
        'Ghi chú: Thái Quang Sơn và Mai Tiểu Bảo đảm nhận 65–70% các task '
        'cốt lõi liên quan đến thiết kế kiến trúc, lập trình backend C#, '
        'xây dựng animation engine và tích hợp hệ thống. Huỳnh Lê Minh Thư '
        'và Trần Viết Tâm Phúc phụ trách các task khảo sát, xây dựng giao '
        'diện (Vue.js Components cơ bản), viết tài liệu và kiểm thử.',
        italic=True)

    add_page_break(doc)


def run() -> None:
    filepath = get_output_path()
    doc = load_document(filepath)
    print('[Phase 2] Đang thêm: Phần 2 – Khảo sát...')
    build_phan2(doc)
    save_document(doc, filepath)
    print('[Phase 2] HOÀN THÀNH.\n')


if __name__ == '__main__':
    run()
