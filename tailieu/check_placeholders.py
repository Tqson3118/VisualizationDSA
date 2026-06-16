import docx
import os
import sys
import io

# Fix Unicode output trên Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def check_placeholders():
    filepath = r"d:\FPT\Hihi\document\PRO2192_Report.docx"
    if not os.path.exists(filepath):
        print("❌ File báo cáo đích không tồn tại!")
        return

    doc = docx.Document(filepath)
    placeholders = []
    
    # Check paragraphs
    for p in doc.paragraphs:
        if "PLACEHOLDER" in p.text:
            placeholders.append(p.text)
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "PLACEHOLDER" in p.text:
                        placeholders.append(p.text)

    if placeholders:
        print(f"❌ Phát hiện {len(placeholders)} placeholder(s):")
        for pl in placeholders:
            print(f"   - {pl}")
    else:
        print("✅ Tuyệt vời! Tài liệu hoàn toàn sạch sẽ, không còn placeholder ảnh nào.")

if __name__ == '__main__':
    check_placeholders()
