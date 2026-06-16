import docx
import sys
import io

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = docx.Document(r"d:\FPT\Hihi\tailieu\PRO2192_Report_Goc.docx")

print("=== PARAGRAPHS IN ORIGINAL DOCUMENT ===")
for i, p in enumerate(doc.paragraphs):
    if len(p.text.strip()) > 0:
        print(f"P[{i}]: {p.text[:100]} (Style: {p.style.name})")

print("\n=== TABLES IN ORIGINAL DOCUMENT ===")
for i, table in enumerate(doc.tables):
    print(f"\nTable[{i}] has {len(table.rows)} rows, {len(table.columns)} columns:")
    for r_idx, row in enumerate(table.rows[:5]): # Xem 5 hàng đầu tiên
        row_text = [cell.text.strip().replace('\n', ' ')[:30] for cell in row.cells]
        print(f"  Row[{r_idx}]: {row_text}")
